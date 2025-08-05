# create_sample_doc.py
from docx import Document

def create_sample_docx():
    """Create a sample DOCX file for testing"""
    doc = Document()
    
    # Add a title
    doc.add_heading('Sample Policy Document', 0)
    
    # Add some content
    doc.add_paragraph('This is a sample policy document created for testing purposes.')
    doc.add_paragraph('It contains multiple paragraphs to simulate a real document.')
    
    doc.add_heading('Section 1: Introduction', level=1)
    doc.add_paragraph('This section introduces the key concepts and objectives of the policy.')
    doc.add_paragraph('The policy aims to establish clear guidelines for document processing.')
    
    doc.add_heading('Section 2: Implementation', level=1)
    doc.add_paragraph('This section outlines the implementation details.')
    doc.add_paragraph('Key steps include: document parsing, text extraction, and data processing.')
    
    doc.add_heading('Section 3: Conclusion', level=1)
    doc.add_paragraph('In conclusion, this policy provides a framework for effective document management.')
    
    # Save the document
    doc.save('sample_policy.docx')
    print("✅ Created sample_policy.docx for testing")

if __name__ == "__main__":
    create_sample_docx()
